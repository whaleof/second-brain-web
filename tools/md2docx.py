#!/usr/bin/env python3
"""Convert 工作台问题汇总.md to a styled Word document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = Path(r"G:\_06_项目代码\工作台\workspace\docs\工作台问题汇总.md")
OUT_PATH = Path(r"G:\_06_项目代码\工作台\workspace\docs\工作台问题汇总.docx")


def parse_bold(text: str, run):
    """Split text by **bold** markers and apply bold to those segments."""
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for i, part in enumerate(parts):
        r = run._element
        # We need to add multiple runs; the initial run is already there.
        # Easier: clear run text and rebuild on the paragraph level.
        pass


def add_rich_text(paragraph, text: str):
    """Add text with **bold** segments to a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
        run.font.size = Pt(10.5)
        run.font.name = "Microsoft YaHei"


def shade_paragraph(paragraph, fill: str):
    """Add a background shade to a paragraph via w:shd."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def main():
    doc = Document()

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)

    # Heading styles
    for lvl, size, color in [(1, 18, RGBColor(0x4A, 0x4A, 0x4A)),
                              (2, 14, RGBColor(0x6B, 0x4C, 0x9A)),
                              (3, 12, RGBColor(0x33, 0x33, 0x33))]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Microsoft YaHei"
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True

    content = MD_PATH.read_text(encoding="utf-8")
    lines = content.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            txt = stripped[3:]
            m = re.match(r"^(.+?)（(.+?)）$", txt)
            if m:
                # Stage band: small title + sub, shaded background, no number
                p = doc.add_heading(level=3)
                run = p.add_run(m.group(1))
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x4A, 0x3B, 0x6B)
                sub = p.add_run("  " + m.group(2))
                sub.font.size = Pt(9)
                sub.font.color.rgb = RGBColor(0x84, 0x73, 0xA8)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                shade_paragraph(p, "F1ECF9")
            else:
                doc.add_heading(txt, level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("> "):
            # Quote block as italic paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.font.size = Pt(10)
        elif stripped.startswith("| ") and "|" in stripped[2:]:
            # Simple table handling for the panorama table
            # Collect table rows until empty line
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                # Skip separator rows
                if not all(set(c) <= set("-:| ") for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text
            continue
        elif re.match(r"^[-*]\s+", stripped):
            # Bullet list item (could be nested)
            level = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + level * 0.15)
            add_rich_text(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped == "":
            # Add small spacing instead of empty paragraphs
            pass
        elif stripped.startswith("---"):
            # Horizontal rule: add a tiny spacer
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_rich_text(p, stripped)

        i += 1

    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    doc.save(OUT_PATH)
    print(f"Saved Word document: {OUT_PATH}")


if __name__ == "__main__":
    main()
